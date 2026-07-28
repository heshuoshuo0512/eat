#!/usr/bin/env python3
"""Normalize supplied campus dining files into an auditable JSON bundle."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import openpyxl
from docx import Document


DATA_VERSION = "west-main-2026-07-27-v1"
PRICE_RE = re.compile(r"(?P<amount>\d+(?:\.\d+)?)\s*元")
STORE_RE = re.compile(r"^店铺\s*\d+\s*[：:]\s*(.+)$")
DOC_STORE_RE = re.compile(r"^[一二三四五六七八九十百]+[、.．]\s*(.+)$")
NUMBER_PREFIX_RE = re.compile(r"^(?:\d+[.、）)]\s*|[（(]\d+[）)]\s*)")
TIER_RE = re.compile(
    r"^[\-•·]?\s*(?:(\d+(?:\.\d+)?)\s*元(?:档|区)?|([一二三四五六七八九十])元区)"
    r"(?:\s*[：:]\s*(.*)|\s*)$"
)
CHINESE_AMOUNTS = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}


@dataclass(frozen=True)
class SourceSpec:
    filename: str
    area_id: str
    area_name: str
    floor: str
    format: str
    parent_id: str = "campus-main"
    profile: str = "default"


@dataclass(frozen=True)
class VenueSpec:
    id: str
    name: str
    location: str
    description: str


@dataclass(frozen=True)
class CatalogSpec:
    key: str
    title: str
    batch_id: str
    data_version: str
    venues: tuple[VenueSpec, ...]
    sources: tuple[SourceSpec, ...]


WEST_SOURCES = (
    SourceSpec("心怡餐厅(4).xlsx", "west-xinyi", "心怡餐厅", "未标注", "xlsx"),
    SourceSpec("民族餐厅(2).docx", "west-minzu", "民族餐厅", "未标注", "docx"),
    SourceSpec("二楼东厅(1).docx", "west-floor2-east", "二楼东厅", "2F东", "docx"),
    SourceSpec("二楼西.txt", "west-xijinjia", "禧进甲餐厅", "2F西", "txt"),
    SourceSpec("[只读]三楼西(2).docx", "west-darongshu", "大榕树餐厅", "3F西", "docx"),
    SourceSpec("三楼东.txt", "west-floor3-east", "三楼东厅", "3F东", "txt"),
)

EAST_SOURCES = (
    SourceSpec("东区燕鸣湖食堂.md", "east-yanminghu-1f", "燕鸣湖餐厅一楼", "1F", "txt", "east-zone", "yanminghu_floor1"),
    SourceSpec("东校区二楼.docx", "east-yanminghu-2f", "燕鸣湖餐厅二楼", "2F", "docx", "east-zone", "yanminghu_floor2"),
    SourceSpec("东区东大活.md", "east-dongdahuo", "东大活", "未标注", "txt", "east-dongdahuo", "dongdahuo"),
    SourceSpec("广源超市(1).docx", "east-guangyuan", "广源超市", "未标注", "docx", "east-guangyuan", "guangyuan"),
)

WEST_VENUE = VenueSpec("campus-main", "西区大食堂", "西区", "西区大食堂真实目录，包含六个次级餐厅。")
YANMINGHU_VENUE = VenueSpec("east-zone", "东区燕鸣湖", "东区", "东区燕鸣湖真实目录，包含一楼与二楼。")
DONGDAHUO_VENUE = VenueSpec("east-dongdahuo", "东大活", "东区", "东大活真实档口目录。")
GUANGYUAN_VENUE = VenueSpec("east-guangyuan", "广源超市", "东区", "广源超市真实档口与商品目录。")

CATALOGS = {
    "west-main": CatalogSpec(
        "west-main",
        "西区大食堂真实菜单清洗报告",
        "real-catalog-west-main-2026-07-27",
        "west-main-2026-07-27-v1",
        (WEST_VENUE,),
        WEST_SOURCES,
    ),
    "campus": CatalogSpec(
        "campus",
        "校园联合真实目录清洗报告",
        "real-catalog-campus-2026-07-27-v2",
        "campus-catalog-2026-07-27-v2",
        (WEST_VENUE, YANMINGHU_VENUE, DONGDAHUO_VENUE, GUANGYUAN_VENUE),
        WEST_SOURCES + EAST_SOURCES,
    ),
}

# Backwards-compatible module-level view used by existing ad-hoc analysis helpers.
SOURCES = list(WEST_SOURCES)


@dataclass
class ParseState:
    spec: SourceSpec
    source_name: str
    source_hash: str
    stall_name: str = ""
    stall_id: str = ""
    tier_amount: float | None = None
    variant_labels: list[str] = field(default_factory=list)
    group_dishes: list[dict[str, Any]] = field(default_factory=list)
    section_suffix: str = ""
    group_suffix: str = ""


def normalize_text(value: Any) -> str:
    text = unicodedata.normalize("NFKC", str(value or ""))
    text = text.replace("‑", "-").replace("–", "-").replace("—", "-").replace("﹣", "-")
    text = re.sub(r"[\u00a0\u2007\u202f\t]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def stable_id(prefix: str, *parts: str) -> str:
    raw = "\x1f".join(normalize_text(part).lower() for part in parts)
    return f"{prefix}-{hashlib.sha256(raw.encode('utf-8')).hexdigest()[:14]}"


def amount_text(value: float) -> str:
    return str(int(value)) if float(value).is_integer() else (f"{value:.2f}".rstrip("0").rstrip("."))


def split_names(value: str) -> list[str]:
    value = normalize_text(value).strip("，,、;；。 ")
    if not value:
        return []
    parts = split_top_level(value)
    return [normalize_text(item).strip("-:：。 ") for item in parts if normalize_text(item).strip("-:：。 ")]


def split_top_level(value: str) -> list[str]:
    parts: list[str] = []
    current: list[str] = []
    depth = 0
    for char in normalize_text(value):
        if char in "（(【[":
            depth += 1
        elif char in "）)】]" and depth:
            depth -= 1
        if char in "、，,；;" and depth == 0:
            part = normalize_text("".join(current)).strip()
            if part:
                parts.append(part)
            current = []
        else:
            current.append(char)
    tail = normalize_text("".join(current)).strip()
    if tail:
        parts.append(tail)
    return parts


def clean_name(value: str) -> str:
    value = NUMBER_PREFIX_RE.sub("", normalize_text(value))
    value = re.sub(r"^[\-•·]+\s*", "", value)
    return value.strip("-‐:：;；。 ")


def aliases_for(name: str) -> list[str]:
    variants = {normalize_text(name)}
    replacements = [("肉沫", "肉末"), ("蒜台", "蒜苔"), ("千页豆腐", "千叶豆腐"), ("西红柿", "番茄")]
    for left, right in replacements:
        for value in list(variants):
            if left in value:
                variants.add(value.replace(left, right))
            if right in value:
                variants.add(value.replace(right, left))
    compact = re.sub(r"[\s\p{P}]", "", name) if False else re.sub(r"[\s\-_/&+()（）·]", "", name)
    if compact:
        variants.add(compact)
    variants.discard(name)
    return sorted(value for value in variants if value and value != name)


def semantic_labels(name: str, context: str = "") -> list[str]:
    text = f"{name} {context}"
    rules = [
        (r"粥|豆浆|包子|烧麦|馅饼|鸡蛋饼|早餐", "早餐"),
        (r"米饭|盖饭|拌饭|炒饭|焖饭|滑蛋饭|烤盘饭", "米饭套餐"),
        (r"面|刀削|板面|拉面|馄饨|饺子|米线|米粉|河粉|螺蛳粉", "粉面主食"),
        (r"麻辣|香辣|辣子|酸辣|剁椒|小米辣", "辣味"),
        (r"清汤|清淡|菌汤", "清淡"),
        (r"汉堡|鸡排|炸鸡|薯条|鸡米花", "西式快餐"),
        (r"麻辣烫|麻辣拌|香锅|小火锅", "自选称重"),
        (r"奶茶|果茶|咖啡|酸梅汤|柠檬水|水吧", "饮品"),
        (r"水饺|锅贴|馄饨|云吞", "饺子馄饨"),
        (r"鸡胸|牛肉|鱼|虾|鸡蛋|豆腐", "蛋白质菜品"),
    ]
    return [label for pattern, label in rules if re.search(pattern, text)]


def classify_non_dish(text: str) -> str | None:
    stripped = clean_name(text)
    if not stripped:
        return "empty"
    if re.search(r"^(?:口味|可选口味|规格|计价|单价|售价|价格|加料|加价|附加|加菜|通用加料|单点加料|可单点加料|炒制加料|卤味|福利|配套服务)[：:]?", stripped):
        return "metadata_or_modifier"
    if re.search(r"(?:系列|类别|分类|类|汤面|拌面|炒面|炒饭|炒粉|炒饼|炒河粉|锅贴|甜品|单品|快餐|早餐|粥品|水煮肉片|加料区|推荐爆款单品)[：:]?$", stripped):
        return "section_heading"
    if stripped in {
        "加料", "加菜", "自选菜", "汤面", "拌面", "炒面", "炒方便面", "滑蛋饭", "锅贴", "特色成品",
        "韩式炸鸡", "套餐饭", "馅饼+粥品", "特色", "冷面", "招牌必点", "手工馄饨", "面皮", "其他饮品",
        "自选香锅 / 麻辣烫 / 麻辣拌", "掉渣饼", "夹饼自选加料", "剁椒面", "水煮肉片", "称重炸物(500g)",
    }:
        return "section_heading"
    if re.search(r"^(?:新增.*味|全店免费|以上均可|米饭免费|.*免费续(?:饭|面|粉)|.*赠送.*)$", stripped):
        return "descriptive_metadata"
    return None


def named_price_items(text: str) -> list[tuple[str, str]]:
    """Return distinct name/price fragments when each top-level segment names one product."""
    segments = split_top_level(text)
    if len(segments) < 2:
        return []
    items: list[tuple[str, str]] = []
    for segment in segments:
        matches = list(PRICE_RE.finditer(segment))
        if len(matches) != 1:
            return []
        name = line_name_before_price(segment)
        if not name or re.fullmatch(r"(?:会员价|统一|售价|单价|价格)", name):
            return []
        items.append((name, segment))
    return items


def ambiguous_name_reason(name: str) -> str | None:
    normalized = clean_name(name)
    if not normalized:
        return "ambiguous_product_name"
    if re.fullmatch(r"(?:单价|售价|计价|价格|荤素统一售价|基础套餐|第一组小锅|第二组小锅)", normalized):
        return "pricing_rule_without_product"
    if re.fullmatch(r"\d+(?:\s*[-~至]\s*\d+)?\s*人份", normalized):
        return "serving_tier_without_product"
    if re.match(r"^(?:第[一二三四五六七八九十]+组小锅|乳酸菌系列|粥品|双拼饭系列)\s*[（(]?.*$", normalized):
        return "group_price_without_products"
    if re.match(r"^(?:以上均可|全店免费|新增.*味|米饭免费)", normalized):
        return "descriptive_metadata"
    return None


def pricing_for(text: str, inherited: float | None = None, variant_labels: list[str] | None = None,
                structured_units: bool = False) -> dict[str, Any] | None:
    normalized = normalize_text(text)
    amounts = [float(match.group("amount")) for match in PRICE_RE.finditer(normalized)]
    if not amounts and inherited is None:
        return None
    if not amounts:
        amounts = [float(inherited)]
    weight_quantity = r"\d+(?:\.\d+)?|[一二三四五六七八九十]+" if structured_units else r"\d+(?:\.\d+)?"
    weight_units = r"克|g|kg|千克|斤|两" if structured_units else r"克|g|kg|千克|斤"
    per_weight = re.search(rf"元\s*/\s*(?:({weight_quantity})\s*)?({weight_units})", normalized, re.I)
    per_unit = re.search(r"元\s*/\s*(?:(\d+(?:\.\d+)?)\s*)?(个|只|笼|份)", normalized)
    per_person = re.search(r"元\s*/\s*(?:每)?位|元\s*每位", normalized)
    mode = "fixed"
    unit = "份"
    base_quantity = None
    budget_comparable = True
    if structured_units and len(amounts) > 1:
        mode = "variants"
    elif per_weight:
        mode = "per_weight"
        raw_quantity = per_weight.group(1) or "1"
        base_quantity = float(raw_quantity) if re.fullmatch(r"\d+(?:\.\d+)?", raw_quantity) else float(CHINESE_AMOUNTS.get(raw_quantity, 1))
        unit = normalize_text(per_weight.group(2)).lower().replace("g", "克")
        budget_comparable = False
    elif structured_units and per_unit:
        mode = "per_unit"
        base_quantity = float(per_unit.group(1) or 1)
        unit = normalize_text(per_unit.group(2))
    elif per_person:
        mode = "per_person"
        unit = "位"
    elif len(amounts) > 1:
        mode = "variants"
    elif inherited is not None and not PRICE_RE.search(normalized):
        mode = "tiered"
    labels = variant_labels or []
    variants = []
    if len(amounts) > 1:
        for index, amount in enumerate(amounts):
            label = labels[index] if index < len(labels) else f"规格{index + 1}"
            variants.append({"id": f"variant-{index + 1}", "label": label, "amount": amount, "quantity": None, "unit": unit})
    display = re.sub(r"\s+", "", normalized)
    price_fragment = re.search(r"\d+(?:\.\d+)?\s*元(?:\s*/\s*[^，,；;）)]*)?", normalized)
    if mode in {"fixed", "tiered"}:
        display = f"{amount_text(amounts[0])}元"
    elif mode == "per_person":
        display = f"{amount_text(amounts[0])}元/位"
    elif mode == "per_weight" and per_weight:
        quantity = amount_text(base_quantity or 1)
        display = f"{amount_text(amounts[0])}元/{quantity}{unit}"
    elif mode == "per_unit" and per_unit:
        quantity = amount_text(base_quantity or 1)
        display = f"{amount_text(amounts[0])}元/{quantity}{unit}"
    elif mode == "variants":
        display = f"{amount_text(min(amounts))}-{amount_text(max(amounts))}元"
    return {
        "mode": mode,
        "display": display,
        "baseAmount": amounts[0],
        "baseQuantity": base_quantity,
        "unit": unit,
        "minAmount": min(amounts),
        "maxAmount": max(amounts),
        "budgetComparable": budget_comparable,
        "variants": variants,
        "modifiers": [],
        "raw": normalized,
    }


def line_name_before_price(text: str) -> str:
    value = normalize_text(text)
    match = PRICE_RE.search(value)
    if not match:
        return clean_name(value)
    before = value[: match.start()].rstrip("-:：/ ")
    if "：" in before or ":" in before:
        before = re.split(r"[：:]", before)[-1]
    return clean_name(before)


def add_audit(audits: list[dict[str, Any]], state: ParseState, locator: str, entity_type: str,
              status: str, raw_text: str, normalized: dict[str, Any] | None = None,
              issues: list[str] | None = None, entity_id: str | None = None) -> None:
    audits.append({
        "id": stable_id("import-row", state.source_hash, locator, entity_type, entity_id or status),
        "sourceHash": state.source_hash,
        "sourceName": state.source_name,
        "sourceLocator": locator,
        "entityType": entity_type,
        "entityId": entity_id,
        "status": status,
        "rawText": normalize_text(raw_text),
        "normalized": normalized or {},
        "issues": issues or [],
    })


def make_stall(state: ParseState, name: str, stalls: list[dict[str, Any]], locator: str,
               audits: list[dict[str, Any]], inferred: bool = False,
               extra_issues: list[str] | None = None) -> None:
    canonical = clean_name(name).replace(". ", "")
    state.stall_name = canonical
    state.stall_id = stable_id("stall", state.spec.area_id, canonical)
    state.tier_amount = None
    state.variant_labels = []
    state.group_dishes = []
    state.section_suffix = ""
    state.group_suffix = ""
    if not any(item["id"] == state.stall_id for item in stalls):
        aliases = ["益禾堂"] if canonical == "益和堂" else []
        generic_name = state.spec.profile == "guangyuan" and canonical == "超市"
        stall = {
            "id": state.stall_id,
            "canteenId": state.spec.area_id,
            "floor": state.spec.floor,
            "name": canonical,
            "category": "名称待核验" if generic_name else "待核验",
            "rating": 0,
            "avgPrice": 0,
            "open": False,
            "description": "原始名称为泛化店名，名称待核验。" if generic_name else ("来源菜单未提供档口说明。" if not inferred else "原始资料未提供档口名，待运营核验。"),
            "aliases": aliases,
        }
        stalls.append(stall)
        issues = list(extra_issues or [])
        if inferred:
            issues.append("inferred_stall_name")
        if generic_name:
            issues.append("generic_stall_name")
        if aliases:
            issues.append("search_alias_added")
        add_audit(audits, state, locator, "stall", "accepted", name, stall, issues, state.stall_id)


def make_dish(state: ParseState, name: str, pricing: dict[str, Any], locator: str, raw_text: str,
              dishes: list[dict[str, Any]], audits: list[dict[str, Any]], context: str = "") -> dict[str, Any] | None:
    name = clean_name(name)
    if not name or len(name) > 100 or not state.stall_id:
        add_audit(audits, state, locator, "dish", "review_required", raw_text, issues=["invalid_name_or_missing_stall"])
        return None
    duplicate = next((dish for dish in dishes if dish["stallId"] == state.stall_id and dish["name"] == name and dish["priceDisplay"] == pricing["display"]), None)
    source = {"name": state.source_name, "sha256": state.source_hash, "locator": locator, "rawText": normalize_text(raw_text)}
    if duplicate:
        duplicate["sourceRef"]["sources"].append(source)
        add_audit(audits, state, locator, "dish", "excluded", raw_text, {"duplicateOf": duplicate["id"]}, ["duplicate_same_stall_name_price"], duplicate["id"])
        return duplicate
    dish_id = stable_id("dish", state.spec.area_id, state.stall_id, name, pricing["display"])
    labels = semantic_labels(name, f"{context} {state.stall_name} {state.spec.area_name}")
    dish = {
        "id": dish_id,
        "stallId": state.stall_id,
        "name": name,
        "price": pricing["minAmount"],
        "priceDisplay": pricing["display"],
        "pricingMode": pricing["mode"],
        "pricing": pricing,
        "taste": "待核验",
        "cuisine": "待核验",
        "ingredients": [],
        "seasonings": [],
        "additives": [],
        "tags": [],
        "aliases": aliases_for(name),
        "semanticLabels": labels,
        "halal": False,
        "mealTypes": [],
        "nutrition": {"calories": 0, "protein": 0, "fat": 0, "carbs": 0},
        "fiber": 0,
        "sodium": 0,
        "sugar": 0,
        "calcium": 0,
        "iron": 0,
        "rating": 0,
        "reviewCount": 0,
        "sales": 0,
        "image": "",
        "imageUrl": "",
        "description": f"{state.spec.area_name} · {state.stall_name}菜单目录；今日供应待确认。",
        "status": "active",
        "allergens": [],
        "safetyDeclarations": [{
            "allergenCode": "*", "status": "unknown", "source": "menu_document",
            "verifiedBy": None, "verifiedAt": None, "expiresAt": None, "dataVersion": DATA_VERSION,
        }],
        "dietaryLabels": [],
        "factStatus": {"nutrition": "unknown", "recipe": "unknown", "halal": "unknown", "dietary": "unknown", "spice": "unknown"},
        "spiceLevel": None,
        "factSource": "menu_document",
        "factVerifiedAt": None,
        "factExpiresAt": None,
        "dataVersion": DATA_VERSION,
        "synthetic": False,
        "sourceRef": {"batchId": "", "dataVersion": DATA_VERSION, "sources": [source]},
    }
    dishes.append(dish)
    state.group_dishes.append(dish)
    add_audit(audits, state, locator, "dish", "accepted", raw_text, dish, [], dish_id)
    return dish


def apply_modifiers(state: ParseState, text: str, locator: str, audits: list[dict[str, Any]]) -> bool:
    prefix_removed = re.sub(
        r"^[\-•·]?\s*(?:加料|加价单品|附加|加菜|通用加料|单点加料|可单点加料|炒制加料|卤味)\s*[：:]\s*",
        "",
        normalize_text(text),
    )
    parts = split_top_level(prefix_removed)
    modifiers = []
    for part in parts:
        amounts = [float(item.group("amount")) for item in PRICE_RE.finditer(part)]
        name = line_name_before_price(part)
        if amounts and name and not re.search(r"^(?:加料|加价单品|附加|加菜|卤味|通用加料)$", name):
            modifiers.append({"label": name, "amount": amounts[0]})
    if not modifiers:
        return False
    for dish in state.group_dishes:
        existing = {(item["label"], item["amount"]) for item in dish["pricing"]["modifiers"]}
        dish["pricing"]["modifiers"].extend(item for item in modifiers if (item["label"], item["amount"]) not in existing)
    add_audit(audits, state, locator, "modifier", "excluded", text, {"modifiers": modifiers, "attachedDishCount": len(state.group_dishes)}, ["stored_as_modifier_not_dish"])
    return True


def with_group_suffix(name: str, suffix: str) -> str:
    cleaned = clean_name(name)
    if not suffix or cleaned.endswith(("饭", "面", "粉", "饺", "锅贴")):
        return cleaned
    return f"{cleaned}{suffix}"


def process_yanminghu_floor1_special(state: ParseState, text: str, locator: str,
                                     dishes: list[dict[str, Any]], audits: list[dict[str, Any]]) -> bool:
    colon_variants = re.match(r"^(.+?)[：:]\s*(.+)$", text)
    if colon_variants and len(PRICE_RE.findall(colon_variants.group(2))) >= 2:
        right = colon_variants.group(2)
        labels = re.findall(r"(一份|半份|小份|中份|大份|\d+个)\s*\d", right)
        pricing = pricing_for(right, variant_labels=labels, structured_units=True)
        if pricing and pricing["mode"] == "variants":
            make_dish(state, clean_name(colon_variants.group(1)), pricing, locator, text, dishes, audits)
            state.tier_amount = None
            state.group_suffix = ""
            return True
    price_first = re.match(r"^(\d+(?:\.\d+)?)\s*元\s*/\s*(个|份|位)\s*[：:]\s*(.+)$", text)
    if price_first:
        pricing = pricing_for(f"{price_first.group(1)}元/{price_first.group(2)}", structured_units=True)
        names = split_names(price_first.group(3))
        for index, name in enumerate(names, 1):
            make_dish(state, name, pricing.copy(), f"{locator}#{index}", text, dishes, audits)
        state.tier_amount = None
        state.group_suffix = ""
        return True

    group_price = re.match(
        r"^(.+?)(?:系列)?[（(].*?(?:全部|统一)\s*(\d+(?:\.\d+)?)\s*元(?:\s*/\s*份)?.*?[）)]$",
        text,
    )
    if group_price:
        group_name = clean_name(group_price.group(1)).removesuffix("系列")
        state.tier_amount = float(group_price.group(2))
        state.group_suffix = group_name if group_name in {"烤肉拌饭", "脆皮鸡饭"} else ""
        state.group_dishes = []
        add_audit(
            audits, state, locator, "section", "excluded", text,
            {"tierAmount": state.tier_amount, "groupName": group_name}, ["group_price_heading"],
        )
        return True
    return False


def normalize_implicit_unit_prices(text: str) -> str:
    normalized = re.sub(r"(?<!元)(\d+(?:\.\d+)?)\s*/\s*(斤|个)", r"\1元/\2", text)
    return re.sub(r"(?<=[斤个])\s+(?=[^，,；;）)])", "、", normalized)


def process_yanminghu_floor2_special(state: ParseState, text: str, locator: str,
                                     dishes: list[dict[str, Any]], audits: list[dict[str, Any]]) -> tuple[bool, str]:
    text = normalize_implicit_unit_prices(text)
    if "收费方式" in text and "9.9元/四两" in text:
        pricing = pricing_for("9.9元/四两", structured_units=True)
        make_dish(state, "自选烤盘饭", pricing, locator, text, dishes, audits, "称重")
        add_audit(audits, state, locator, "metadata", "excluded", text, issues=["pricing_rule_attached_to_catalog_dish"])
        return True, text
    if re.search(r"(?:基础可选口味|全部可选口味|整体规则|^备注[：:]|^优惠[：:])", text):
        add_audit(audits, state, locator, "metadata", "excluded", text, issues=["descriptive_metadata"])
        return True, text
    section = re.match(r"^[（(]?[12][）)]?\s*(手工水饺|锅贴)$", text)
    if section:
        state.section_suffix = "水饺" if "水饺" in section.group(1) else "锅贴"
        state.tier_amount = None
        state.group_dishes = []
        add_audit(audits, state, locator, "section", "excluded", text, {"dishSuffix": state.section_suffix}, ["product_section_heading"])
        return True, text
    tier_heading = re.match(r"^(\d+(?:\.\d+)?)\s*元\s*档位$", text)
    if tier_heading:
        state.tier_amount = float(tier_heading.group(1))
        state.group_dishes = []
        add_audit(audits, state, locator, "section", "excluded", text, {"tierAmount": state.tier_amount}, ["tier_heading"])
        return True, text

    colon = re.match(r"^(.+?)[：:]\s*(.+)$", text)
    if colon and len(PRICE_RE.findall(colon.group(2))) >= 2:
        left = clean_name(colon.group(1))
        right = colon.group(2)
        fragments = split_top_level(right)
        labels = [line_name_before_price(fragment) for fragment in fragments if PRICE_RE.search(fragment)]
        pricing = pricing_for(right, variant_labels=labels, structured_units=True)
        if pricing and pricing["mode"] == "variants":
            names = split_names(left)
            portion_labels = {"一份", "半份", "小份", "中份", "大份", "1个", "3个"}
            if state.section_suffix and labels and all(label in portion_labels for label in labels):
                names = [name if name.endswith(state.section_suffix) else f"{name}{state.section_suffix}" for name in names]
            for index, name in enumerate(names, 1):
                make_dish(state, name, pricing.copy(), f"{locator}#{index}" if len(names) > 1 else locator, text, dishes, audits)
            state.tier_amount = None
            return True, text
    return False, text


def process_text_line(state: ParseState, text: str, locator: str, dishes: list[dict[str, Any]], audits: list[dict[str, Any]]) -> None:
    text = normalize_text(text)
    if not text:
        return
    structured_units = state.spec.parent_id == "east-zone" or state.spec.profile in {"dongdahuo", "guangyuan"}
    if state.spec.profile == "yanminghu_floor1" and process_yanminghu_floor1_special(state, text, locator, dishes, audits):
        return
    if state.spec.profile == "yanminghu_floor2":
        handled, text = process_yanminghu_floor2_special(state, text, locator, dishes, audits)
        if handled:
            return
    price_first_person = re.match(r"^(\d+(?:\.\d+)?)\s*元\s*/\s*(?:位|人)\s*(.+)$", text)
    if price_first_person:
        name = clean_name(price_first_person.group(2))
        pricing = pricing_for(f"{name} {price_first_person.group(1)}元/位", structured_units=structured_units)
        make_dish(state, name, pricing, locator, text, dishes, audits)
        state.tier_amount = None
        return
    tier = TIER_RE.match(text)
    if tier:
        amount = float(tier.group(1)) if tier.group(1) else float(CHINESE_AMOUNTS[tier.group(2)])
        remainder = (tier.group(3) or "").strip()
        state.group_dishes = []
        if remainder:
            state.tier_amount = None
            names = split_names(remainder)
            for index, name in enumerate(names, 1):
                pricing = pricing_for(name, inherited=amount, structured_units=structured_units)
                make_dish(state, name, pricing, f"{locator}#{index}", text, dishes, audits, "价格档")
        else:
            state.tier_amount = amount
            add_audit(audits, state, locator, "section", "excluded", text, {"tierAmount": amount}, ["tier_heading"])
        return
    reason = classify_non_dish(text)
    if reason == "metadata_or_modifier" or re.search(r"^(?:加料|加价|附加|加菜|卤味).*\d", text):
        if apply_modifiers(state, text, locator, audits):
            return
    if reason:
        state.group_dishes = []
        state.tier_amount = None
        state.section_suffix = ""
        state.group_suffix = ""
        if "口味" in text and "/" in text and not PRICE_RE.search(text):
            state.variant_labels = [item.strip() for item in re.split(r"[/、]", re.split(r"[：:]", text)[-1]) if item.strip()]
        entity_type = "metadata" if reason == "descriptive_metadata" else "section"
        add_audit(audits, state, locator, entity_type, "excluded", text, issues=[reason])
        return
    distinct_items = named_price_items(text)
    if distinct_items:
        state.tier_amount = None
        state.group_dishes = []
        for index, (item_name, fragment) in enumerate(distinct_items, 1):
            pricing = pricing_for(fragment, structured_units=structured_units)
            make_dish(state, item_name, pricing, f"{locator}#{index}", text, dishes, audits)
        return
    pricing = pricing_for(text, inherited=state.tier_amount, variant_labels=state.variant_labels, structured_units=structured_units)
    if not pricing:
        if re.search(r"口味|蘸料|免费|续饭|新增|米饭|锅底|配方", text):
            add_audit(audits, state, locator, "metadata", "excluded", text, issues=["descriptive_metadata"])
        else:
            add_audit(audits, state, locator, "dish", "review_required", text, issues=["missing_or_ambiguous_price"])
        return
    name_part = line_name_before_price(text)
    ambiguous_reason = ambiguous_name_reason(name_part)
    if ambiguous_reason:
        add_audit(audits, state, locator, "dish", "review_required", text, issues=[ambiguous_reason])
        return
    names = split_names(name_part) if any(char in name_part for char in "、，,；;") else [name_part]
    for index, name in enumerate(names, 1):
        make_dish(state, with_group_suffix(name, state.group_suffix), pricing.copy(), f"{locator}#{index}" if len(names) > 1 else locator, text, dishes, audits)
    state.variant_labels = []


def decode_text(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Cannot decode {path}")


def parse_text_source(path: Path, spec: SourceSpec, stalls: list[dict[str, Any]], dishes: list[dict[str, Any]], audits: list[dict[str, Any]]) -> None:
    state = ParseState(spec, path.name, sha256_file(path))
    lines = [normalize_text(line) for line in decode_text(path).splitlines() if normalize_text(line)]
    for index, text in enumerate(lines, 1):
        locator = f"line:{index}"
        if index == 1:
            add_audit(audits, state, locator, "area", "excluded", text, {"areaId": spec.area_id}, ["source_title"])
            continue
        store = STORE_RE.match(text)
        if store:
            make_stall(state, store.group(1), stalls, locator, audits)
            continue
        if not state.stall_id:
            add_audit(audits, state, locator, "dish", "review_required", text, issues=["missing_stall_heading"])
            continue
        process_text_line(state, text, locator, dishes, audits)


def parse_guangyuan_source(path: Path, spec: SourceSpec, stalls: list[dict[str, Any]],
                           dishes: list[dict[str, Any]], audits: list[dict[str, Any]]) -> None:
    state = ParseState(spec, path.name, sha256_file(path))
    lines = [normalize_text(paragraph.text) for paragraph in Document(path).paragraphs if normalize_text(paragraph.text)]
    expected_unnumbered_stores = [
        "麦当劳", "橘包包现烤面包", "益和堂", "喜悦冰糖葫芦", "好滋味", "蜜雪冰城", "甜啦啦",
        "幸运咖", "淳菓酸奶", "瑞幸咖啡", "库迪咖啡", "超市", "鸭货",
    ]
    next_store = 0
    for index, text in enumerate(lines, 1):
        locator = f"paragraph:{index}"
        if index == 1:
            add_audit(audits, state, locator, "area", "excluded", text, {"areaId": spec.area_id}, ["source_title"])
            continue
        numbered = DOC_STORE_RE.match(text)
        if numbered:
            make_stall(state, numbered.group(1), stalls, locator, audits)
            continue
        if next_store < len(expected_unnumbered_stores) and text == expected_unnumbered_stores[next_store]:
            make_stall(state, text, stalls, locator, audits, extra_issues=["unnumbered_store_order_inferred"])
            next_store += 1
            continue
        if state.stall_name in {"喜悦冰糖葫芦", "好滋味"} and not PRICE_RE.search(text):
            for item_index, name in enumerate(split_names(text), 1):
                add_audit(
                    audits, state, f"{locator}#{item_index}", "dish", "review_required", name,
                    {"name": name, "stallId": state.stall_id}, ["missing_price"],
                )
            continue
        if state.stall_name == "鸭货" and PRICE_RE.search(text):
            pricing = pricing_for(text, structured_units=True)
            make_dish(state, line_name_before_price(text), pricing, locator, text, dishes, audits)
            continue
        add_audit(audits, state, locator, "row", "review_required", text, issues=["unexpected_guangyuan_structure"])


def parse_docx_source(path: Path, spec: SourceSpec, stalls: list[dict[str, Any]], dishes: list[dict[str, Any]], audits: list[dict[str, Any]]) -> None:
    if spec.profile == "guangyuan":
        parse_guangyuan_source(path, spec, stalls, dishes, audits)
        return
    state = ParseState(spec, path.name, sha256_file(path))
    lines = [normalize_text(paragraph.text) for paragraph in Document(path).paragraphs if normalize_text(paragraph.text)]
    if spec.area_id == "west-minzu":
        make_stall(state, "民族餐厅综合档口", stalls, "paragraph:1", audits, inferred=True)
    elif spec.profile == "yanminghu_floor2" and len(lines) > 1:
        make_stall(state, lines[1], stalls, "paragraph:2", audits)
    for index, text in enumerate(lines, 1):
        locator = f"paragraph:{index}"
        if index == 1:
            add_audit(audits, state, locator, "area", "excluded", text, {"areaId": spec.area_id}, ["source_title"])
            continue
        if spec.profile == "yanminghu_floor2" and index == 2:
            continue
        heading = DOC_STORE_RE.match(text)
        if heading and spec.area_id != "west-minzu":
            make_stall(state, heading.group(1), stalls, locator, audits)
            continue
        if not state.stall_id:
            add_audit(audits, state, locator, "dish", "review_required", text, issues=["missing_stall_heading"])
            continue
        process_text_line(state, text, locator, dishes, audits)


def parse_xlsx_source(path: Path, spec: SourceSpec, stalls: list[dict[str, Any]], dishes: list[dict[str, Any]], audits: list[dict[str, Any]]) -> None:
    state = ParseState(spec, path.name, sha256_file(path))
    sheet = openpyxl.load_workbook(path, data_only=True).active
    block_stalls: dict[int, str] = {}
    for merged in sheet.merged_cells.ranges:
        value = normalize_text(sheet.cell(merged.min_row, merged.min_col).value)
        if merged.min_col == 2 and merged.max_row > merged.min_row and merged.min_row > 1 and value:
            canonical = value.replace("心仪", "心怡")
            for row in range(merged.min_row, merged.max_row + 1):
                block_stalls[row] = canonical
    last_stall = ""
    current_tier: float | None = None
    for row_index in range(1, sheet.max_row + 1):
        values = [(cell.column, normalize_text(cell.value)) for cell in sheet[row_index] if normalize_text(cell.value)]
        if not values:
            continue
        locator = f"{sheet.title}!row:{row_index}"
        if row_index <= 2:
            add_audit(audits, state, locator, "area", "excluded", " | ".join(value for _, value in values), {"areaId": spec.area_id}, ["source_title"])
            continue
        stall_name = block_stalls.get(row_index) or last_stall
        if block_stalls.get(row_index) and block_stalls[row_index] != last_stall:
            make_stall(state, block_stalls[row_index], stalls, locator, audits)
            last_stall = block_stalls[row_index]
            current_tier = None
        if not stall_name or not state.stall_id:
            add_audit(audits, state, locator, "dish", "review_required", " | ".join(value for _, value in values), issues=["missing_stall_heading"])
            continue
        row_text = " | ".join(value for column, value in values if column > 1 and value != stall_name)
        cells = [(column, value) for column, value in values if column > 1 and value != stall_name]
        tier_cell = next((value for _, value in cells if TIER_RE.match(value)), None)
        if tier_cell:
            tier_match = TIER_RE.match(tier_cell)
            current_tier = float(tier_match.group(1)) if tier_match.group(1) else float(CHINESE_AMOUNTS[tier_match.group(2)])
            state.tier_amount = current_tier
            state.group_dishes = []
        consumed: set[int] = set()
        made = 0
        for idx, (column, value) in enumerate(cells):
            if idx in consumed or PRICE_RE.search(value) or TIER_RE.match(value) or classify_non_dish(value):
                continue
            next_entry = cells[idx + 1] if idx + 1 < len(cells) else None
            if next_entry and PRICE_RE.search(next_entry[1]):
                if "加" in value and state.group_dishes:
                    apply_modifiers(state, f"{value}:{next_entry[1]}", f"{locator}:cell:{column}", audits)
                else:
                    pricing = pricing_for(f"{value}:{next_entry[1]}")
                    if pricing:
                        make_dish(state, value, pricing, f"{locator}:cell:{column}", f"{value} {next_entry[1]}", dishes, audits)
                        made += 1
                consumed.update({idx, idx + 1})
            elif current_tier is not None and column >= 4:
                pricing = pricing_for(value, inherited=current_tier)
                make_dish(state, value, pricing, f"{locator}:cell:{column}", value, dishes, audits, "价格档")
                made += 1
                consumed.add(idx)
        if made == 0:
            reason = "section_or_metadata" if any(classify_non_dish(value) for _, value in cells) or tier_cell else "unparsed_xlsx_row"
            add_audit(audits, state, locator, "row", "excluded" if reason == "section_or_metadata" else "review_required", row_text, issues=[reason])


def build_canteens(catalog: CatalogSpec) -> list[dict[str, Any]]:
    canteens = []
    venue_by_id = {venue.id: venue for venue in catalog.venues}
    for venue in catalog.venues:
        canteens.append({
            "id": venue.id, "name": venue.name, "location": venue.location, "hours": "待核验",
            "crowdLevel": 0, "tags": ["真实目录", "供应待确认"], "description": venue.description,
            "parentId": None, "canteenType": "primary", "imageUrl": "",
        })
    for spec in catalog.sources:
        if spec.area_id in venue_by_id:
            continue
        venue = venue_by_id[spec.parent_id]
        location = f"{venue.name} · {spec.floor}" if spec.floor != "未标注" else f"{venue.name} · 楼层未标注"
        tags = ["真实目录", "供应待确认"]
        if spec.profile in {"dongdahuo", "guangyuan"}:
            tags.append("服务楼")
        canteens.append({
            "id": spec.area_id, "name": spec.area_name, "location": location, "hours": "待核验",
            "crowdLevel": 0, "tags": tags,
            "description": "目录来自2026-07-27提供的原始资料快照，营业、实时供应与菜单时效均待核验。",
            "parentId": spec.parent_id, "canteenType": "sub", "imageUrl": "",
        })
    return canteens


def quality_report(bundle: dict[str, Any], catalog: CatalogSpec) -> str:
    report = bundle["report"]
    lines = [
        f"# {catalog.title}", "", f"- 批次：`{bundle['manifest']['batchId']}`",
        f"- 数据版本：`{catalog.data_version}`", f"- 顶层区域：{report['venueCount']}",
        f"- 来源餐饮区域：{report['areaCount']}", f"- 目录节点：{report['canteenCount']}", f"- 档口：{report['stallCount']}",
        f"- 已接受菜品：{report['dishCount']}", f"- 待核验记录：{report['reviewRequiredCount']}",
        f"- 排除/元数据记录：{report['excludedCount']}", "", "## 计价方式", "",
    ]
    for mode, count in sorted(report["pricingModes"].items()):
        lines.append(f"- `{mode}`：{count}")
    lines += ["", "## 安全边界", "", "- 未提供的配方、营养、过敏原、清真和饮食模式全部保持 unknown。", "- 未创建今日菜单，所有菜品仅作为目录检索。", "- 待核验记录未写入菜品表或向量索引。", ""]
    return "\n".join(lines)


def main() -> None:
    global DATA_VERSION
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--dataset", choices=sorted(CATALOGS), default="west-main")
    parser.add_argument("--batch-id")
    args = parser.parse_args()
    catalog = CATALOGS[args.dataset]
    DATA_VERSION = catalog.data_version
    batch_id = args.batch_id or catalog.batch_id
    source_dir = Path(args.source_dir)
    output_dir = Path(args.output)
    missing = [spec.filename for spec in catalog.sources if not (source_dir / spec.filename).exists()]
    if missing:
        raise SystemExit(f"Missing source files: {', '.join(missing)}")
    stalls: list[dict[str, Any]] = []
    dishes: list[dict[str, Any]] = []
    audits: list[dict[str, Any]] = []
    source_manifest = []
    for spec in catalog.sources:
        path = source_dir / spec.filename
        source_manifest.append({
            "name": path.name, "sha256": sha256_file(path), "bytes": path.stat().st_size,
            "areaId": spec.area_id, "parentVenueId": spec.parent_id,
        })
        if spec.format == "xlsx":
            parse_xlsx_source(path, spec, stalls, dishes, audits)
        elif spec.format == "docx":
            parse_docx_source(path, spec, stalls, dishes, audits)
        else:
            parse_text_source(path, spec, stalls, dishes, audits)
    for dish in dishes:
        dish["sourceRef"]["batchId"] = batch_id
    for stall in stalls:
        prices = [dish["price"] for dish in dishes if dish["stallId"] == stall["id"] and dish["pricing"]["budgetComparable"]]
        stall["avgPrice"] = round(sum(prices) / len(prices), 2) if prices else 0
    pricing_modes: dict[str, int] = {}
    for dish in dishes:
        pricing_modes[dish["pricingMode"]] = pricing_modes.get(dish["pricingMode"], 0) + 1
    report = {
        "venueCount": len(catalog.venues), "areaCount": len(catalog.sources),
        "canteenCount": len(build_canteens(catalog)),
        "stallCount": len(stalls), "dishCount": len(dishes),
        "stallsByParent": {
            venue.id: sum(1 for stall in stalls if next((source.parent_id for source in catalog.sources if source.area_id == stall["canteenId"]), None) == venue.id)
            for venue in catalog.venues
        },
        "acceptedAuditCount": sum(row["status"] == "accepted" for row in audits),
        "reviewRequiredCount": sum(row["status"] == "review_required" for row in audits),
        "excludedCount": sum(row["status"] == "excluded" for row in audits),
        "pricingModes": pricing_modes,
    }
    bundle = {
        "manifest": {
            "batchId": batch_id, "tenantId": "default", "dataVersion": catalog.data_version,
            "createdAt": datetime.now(timezone.utc).isoformat(), "sources": source_manifest,
        },
        "canteens": build_canteens(catalog), "stalls": stalls, "dishes": dishes, "menus": [],
        "importRows": audits, "report": report,
    }
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "catalog.json").write_text(json.dumps(bundle, ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "review-required.json").write_text(json.dumps([row for row in audits if row["status"] == "review_required"], ensure_ascii=False, indent=2), encoding="utf-8")
    (output_dir / "quality-report.md").write_text(quality_report(bundle, catalog), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
